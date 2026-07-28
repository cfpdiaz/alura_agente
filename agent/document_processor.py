"""Procesador de documentos Word para extracción de texto."""
import os
from typing import List
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from .config import Config

class DocumentProcessor:
    """Procesa archivos Word y los prepara para indexación."""
    
    def __init__(self, config: Config):
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extrae texto de un archivo .docx."""
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def process_documents(self, data_dir: str = "data") -> List[LangchainDocument]:
        """Procesa todos los documentos Word en el directorio data/."""
        documents = []
        
        for filename in os.listdir(data_dir):
            if filename.endswith(".docx"):
                file_path = os.path.join(data_dir, filename)
                category_key = filename.replace(".docx", "")
                category_name = self.config.DOCUMENT_MAPPING.get(
                    category_key, category_key.title()
                )
                
                text = self.extract_text_from_docx(file_path)
                chunks = self.text_splitter.split_text(text)
                
                for i, chunk in enumerate(chunks):
                    doc = LangchainDocument(
                        page_content=chunk,
                        metadata={
                            "source": filename,
                            "category": category_name,
                            "chunk_id": i,
                            "file_path": file_path
                        }
                    )
                    documents.append(doc)
        
        return documents
